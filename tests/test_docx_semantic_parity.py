from __future__ import annotations

from pathlib import Path
import hashlib
import io
import subprocess
import sys
import tempfile
import unittest
from unittest import mock
from zipfile import ZipFile

try:
    import tomllib
except ModuleNotFoundError:  # Python 3.10 test extra
    import tomli as tomllib

from docx import Document
from lxml import etree

from scripts import docx_export
from scripts.contracts import EvidenceStatus, RecommendationResult
from scripts.report_model import ReportModel, build_report_model, render_markdown
from tests.test_generate_report_evidence import (
    evidence_snapshot,
    formal_pathway_result,
    pathway_result,
    rank_estimate,
    recommendations,
    student,
)


ROOT = Path(__file__).resolve().parents[1]
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def model(**overrides) -> ReportModel:
    values = {
        "profile": student(secondary_subjects=("化学", "生物")),
        "recommendations": recommendations(),
        "rank": rank_estimate(),
        "pathways": formal_pathway_result(),
        "evidence": evidence_snapshot(),
    }
    values.update(overrides)
    return build_report_model(**values)


def document_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def xml_part(path: Path, name: str):
    with ZipFile(path) as package:
        return etree.fromstring(package.read(name))


class DocxSemanticParityTest(unittest.TestCase):
    def test_export_projects_the_complete_report_model_semantics(self):
        """Catches a renderer that drops evidence, rank, pathway, or action fields."""
        report = model()
        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "report.docx"
            result = docx_export.export_docx(report, output)
            text = document_text(result)

        for literal in (
            "匿名升学规划报告（演示甲省）",
            "化学、生物",
            "证据状态",
            "部分覆盖",
            "检索日期",
            "2026-08-23",
            "仅覆盖 2026",
            "虚构甲大学",
            "4300",
            "school_rank_offset_median_observed_spread",
            "虚构正式专项",
            "合成培养安排",
            "合成转段规则",
            "合成出口说明",
            "合成服务就业说明",
            "合成退出规则",
            "合成费用说明",
            "model-report",
            "documented_rank_delta",
            "下一步行动建议",
            "AI 生成，仅供参考",
        ):
            self.assertIn(literal, text)
        self.assertNotIn("http://", text)
        self.assertNotIn("https://", text)
        self.assertNotIn(str(ROOT), text)
        self.assertGreaterEqual(text.count("AI 生成，仅供参考"), 3)
        self.assertIn("一、输入与证据边界", render_markdown(report))

    def test_optional_and_unusable_sections_degrade_without_proxy_values(self):
        """Catches DOCX-only fallbacks that invent rank or pathway values."""
        empty = RecommendationResult(
            items=(),
            excluded_by_subject_count=2,
            zero_score_excluded_count=1,
            input_years=(2026,),
            usable_years=(),
            verified_rank_coverage=None,
            coverage_status=EvidenceStatus.MASKED,
            empty_reason="unusable_evidence",
            warnings=("屏蔽值未进入计算",),
        )
        report = model(recommendations=empty, rank=None, pathways=None)
        with tempfile.TemporaryDirectory() as temporary:
            output = docx_export.export_docx(report, Path(temporary))
            text = document_text(output)

        self.assertEqual(output.name, "匿名升学规划报告-演示甲省-2026.docx")
        self.assertIn("屏蔽、冲突或不可精确使用", text)
        self.assertIn("不执行校排名折算", text)
        self.assertIn("多元升学数据不足", text)
        self.assertNotIn("− 4000", text)
        self.assertNotIn("-4000", text)

    def test_export_rejects_detached_or_missing_document_capability(self):
        """Catches constructor bypass and silent optional-dependency skips."""
        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "report.docx"
            with self.assertRaisesRegex(TypeError, "ReportModel"):
                docx_export.export_docx({"profile": "detached"}, output)
            with mock.patch.object(docx_export, "Document", None):
                with self.assertRaisesRegex(
                    docx_export.DocumentDependencyError, "documents"
                ):
                    docx_export.export_docx(model(), output)
                stderr = io.StringIO()
                with mock.patch("sys.stderr", stderr):
                    self.assertEqual(docx_export.main([]), 2)
                self.assertIn("python-docx", stderr.getvalue())

    def test_cli_repeated_secondary_subject_reaches_docx(self):
        """Catches a parser that accepts repeated subjects but drops them downstream."""
        with tempfile.TemporaryDirectory() as temporary:
            output = Path(temporary) / "cli.docx"
            command = [
                sys.executable,
                str(ROOT / "scripts" / "docx_export.py"),
                "--dataset",
                str(ROOT / "tests" / "fixtures" / "provinces" / "demo-312"),
                "--profile",
                str(ROOT / "tests" / "fixtures" / "profiles" / "demo.json"),
                "--evidence",
                str(ROOT / "tests" / "fixtures" / "evidence" / "three-source-consensus"),
                "--secondary-subject",
                "化学",
                "--secondary-subject",
                "生物",
                "--output",
                str(output),
            ]
            completed = subprocess.run(
                command, capture_output=True, text=True, encoding="utf-8"
            )
            self.assertEqual(completed.returncode, 0, completed.stderr)
            self.assertTrue(output.is_file())
            text = document_text(output)

        self.assertIn("化学、生物", text)
        self.assertNotIn("张三", text)

    def test_pending_pathway_preserves_missing_constraints_and_real_details(self):
        """Catches a DOCX renderer that hides why a pathway is not formal."""
        report = model(pathways=pathway_result())
        with tempfile.TemporaryDirectory() as temporary:
            path = docx_export.export_docx(report, Path(temporary))
            text = document_text(path)
        for literal in (
            "待核实",
            "服务期未核实",
            "合成培养安排",
            "合成转段规则",
            "合成出口说明",
            "当前证据未提供",
            "未提供有依据的位次模型",
        ):
            self.assertIn(literal, text)

    def test_documents_extra_is_a_real_runtime_dependency(self):
        """Catches an optional feature whose tests can only skip at import time."""
        payload = tomllib.loads((ROOT / "pyproject.toml").read_text(encoding="utf-8"))
        self.assertIn("python-docx>=1.1,<2", payload["project"]["optional-dependencies"]["documents"])
        self.assertGreaterEqual(tuple(map(int, __import__("docx").__version__.split(".")[:2])), (1, 1))


class DocxStructureTest(unittest.TestCase):
    def export(self, temporary: str) -> Path:
        return docx_export.export_docx(model(), Path(temporary) / "report.docx")

    def test_compact_reference_guide_geometry_styles_and_metadata(self):
        """Catches reliance on Word defaults or personal core properties."""
        with tempfile.TemporaryDirectory() as temporary:
            path = self.export(temporary)
            document = Document(path)
            section = document.sections[0]
            self.assertEqual(section.page_width.twips, 12240)
            self.assertEqual(section.page_height.twips, 15840)
            self.assertEqual(section.top_margin.twips, 1440)
            self.assertEqual(section.right_margin.twips, 1440)
            self.assertEqual(section.bottom_margin.twips, 1440)
            self.assertEqual(section.left_margin.twips, 1440)
            self.assertEqual(section.header_distance.twips, 708)
            self.assertEqual(section.footer_distance.twips, 708)
            self.assertEqual(document.styles["Normal"].font.name, "Calibri")
            self.assertEqual(document.styles["Normal"].font.size.pt, 11)
            self.assertEqual(document.styles["Normal"].paragraph_format.space_after.pt, 6)
            self.assertEqual(document.styles["Normal"].paragraph_format.line_spacing, 1.25)
            for style_name, size, color, before, after in (
                ("Heading 1", 16, "2E74B5", 18, 10),
                ("Heading 2", 13, "2E74B5", 14, 7),
                ("Heading 3", 12, "1F4D78", 10, 5),
            ):
                style = document.styles[style_name]
                self.assertEqual(style.font.size.pt, size)
                self.assertEqual(str(style.font.color.rgb), color)
                self.assertEqual(style.paragraph_format.space_before.pt, before)
                self.assertEqual(style.paragraph_format.space_after.pt, after)
            self.assertEqual(document.core_properties.author, "")
            self.assertEqual(document.core_properties.last_modified_by, "")
            self.assertNotIn("张三", str(document.core_properties.__dict__))

    def test_tables_use_fixed_matching_dxa_geometry_without_fixed_rows(self):
        """Catches percentage/autofit tables, drifting cell widths, and clipped rows."""
        with tempfile.TemporaryDirectory() as temporary:
            root = xml_part(self.export(temporary), "word/document.xml")
        tables = root.xpath(".//w:tbl", namespaces=NS)
        self.assertGreaterEqual(len(tables), 2)
        for table in tables:
            width = table.xpath("./w:tblPr/w:tblW", namespaces=NS)[0]
            indent = table.xpath("./w:tblPr/w:tblInd", namespaces=NS)[0]
            self.assertEqual(width.get(f"{{{NS['w']}}}type"), "dxa")
            self.assertEqual(width.get(f"{{{NS['w']}}}w"), "9360")
            self.assertEqual(indent.get(f"{{{NS['w']}}}type"), "dxa")
            self.assertEqual(indent.get(f"{{{NS['w']}}}w"), "120")
            grid = [
                int(node.get(f"{{{NS['w']}}}w"))
                for node in table.xpath("./w:tblGrid/w:gridCol", namespaces=NS)
            ]
            self.assertEqual(sum(grid), 9360)
            for row in table.xpath("./w:tr", namespaces=NS):
                widths = [
                    int(node.get(f"{{{NS['w']}}}w"))
                    for node in row.xpath("./w:tc/w:tcPr/w:tcW", namespaces=NS)
                ]
                self.assertEqual(widths, grid)
        self.assertFalse(root.xpath(".//w:tblW[@w:type='pct']", namespaces=NS))
        self.assertFalse(root.xpath(".//w:tblLayout[@w:type='autofit']", namespaces=NS))
        self.assertFalse(root.xpath(".//w:trHeight", namespaces=NS))

    def test_lists_use_real_numbering_and_never_fake_markers(self):
        """Catches manual bullets/numbers and wrapped-line misalignment."""
        with tempfile.TemporaryDirectory() as temporary:
            path = self.export(temporary)
            document_xml = xml_part(path, "word/document.xml")
            numbering_xml = xml_part(path, "word/numbering.xml")
        numbered = document_xml.xpath(".//w:p[w:pPr/w:numPr]", namespaces=NS)
        self.assertGreaterEqual(len(numbered), 3)
        for paragraph in document_xml.xpath(".//w:body/w:p", namespaces=NS):
            text = "".join(paragraph.xpath(".//w:t/text()", namespaces=NS)).lstrip()
            self.assertFalse(text.startswith(("•", "·", "- ")))
            self.assertIsNone(__import__("re").match(r"^\d+[.、]\s", text))
        levels = numbering_xml.xpath(".//w:abstractNum/w:lvl[@w:ilvl='0']", namespaces=NS)
        self.assertTrue(
            any(
                level.xpath("./w:numFmt[@w:val='bullet']", namespaces=NS)
                and level.xpath("./w:pPr/w:ind[@w:left='540'][@w:hanging='271']", namespaces=NS)
                for level in levels
            )
        )
        self.assertTrue(
            any(
                level.xpath("./w:numFmt[@w:val='decimal']", namespaces=NS)
                and level.xpath("./w:pPr/w:ind[@w:left='540'][@w:hanging='271']", namespaces=NS)
                for level in levels
            )
        )

    def test_package_is_byte_deterministic_and_contains_no_local_identity(self):
        """Catches timestamps, temp paths, or author identity leaking into the ZIP."""
        with tempfile.TemporaryDirectory() as temporary:
            first = docx_export.export_docx(model(), Path(temporary) / "a.docx")
            second = docx_export.export_docx(model(), Path(temporary) / "b.docx")
            first_bytes = first.read_bytes()
            second_bytes = second.read_bytes()
            self.assertEqual(hashlib.sha256(first_bytes).digest(), hashlib.sha256(second_bytes).digest())
            with ZipFile(first) as package:
                xml = b"\n".join(
                    package.read(name)
                    for name in package.namelist()
                    if name.endswith((".xml", ".rels"))
                ).decode("utf-8", errors="replace")
        # OOXML namespace declarations are themselves HTTP URIs; raw-source
        # URLs are therefore asserted against visible text in the semantic test.
        for forbidden in (str(ROOT), "C:\\Users\\hp", "张三", "13800138000"):
            self.assertNotIn(forbidden, xml)


if __name__ == "__main__":
    unittest.main()
