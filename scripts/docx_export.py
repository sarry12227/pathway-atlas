"""Deterministic DOCX projection of the evidence-aware ``ReportModel``.

The renderer formats already-decided values.  Only the CLI wrapper reads files
and invokes the same validated pipeline used by the Markdown CLI.
"""

from __future__ import annotations

import argparse
from datetime import datetime
import json
import os
from pathlib import Path
import sys
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # Importable without the extra so the CLI can fail clearly.
    Document = None
    WD_ORIENT = WD_ALIGN_VERTICAL = WD_ALIGN_PARAGRAPH = WD_BREAK = None
    OxmlElement = qn = Inches = Pt = RGBColor = None

if __package__:
    from .contracts import CapabilityTier, EvidenceStatus, RecommendationProfile
    from .path_recommend import evaluate_pathways
    from .report_model import ReportModel, StudentProfile, build_report_model
else:  # pragma: no cover - direct script execution
    from contracts import CapabilityTier, EvidenceStatus, RecommendationProfile
    from path_recommend import evaluate_pathways
    from report_model import ReportModel, StudentProfile, build_report_model


CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
HEADING_BLUE = "2E74B5"
HEADING_DARK_BLUE = "1F4D78"
INK_BLUE = "0B2545"
MUTED = "666666"
DISCLAIMER = "AI 生成，仅供参考；不构成录取承诺，最终以当年官方发布为准。"
PUBLIC_DOCX_BASENAME = "anonymous-admission-report.docx"

_STATUS_LABEL = {
    EvidenceStatus.OFFICIAL: "官方",
    EvidenceStatus.CORROBORATED: "多源核验",
    EvidenceStatus.REFERENCE: "多源参考",
    EvidenceStatus.INFERRED: "推断",
    EvidenceStatus.CONFLICT: "冲突",
    EvidenceStatus.MISSING: "缺失",
    EvidenceStatus.MASKED: "屏蔽",
    EvidenceStatus.PARTIAL: "部分覆盖",
}
_TIER_LABEL = {
    CapabilityTier.FULL: "完整档",
    CapabilityTier.STANDARD: "标准档",
    CapabilityTier.OFFLINE: "离线档",
}
_CONFIDENCE_LABEL = {
    EvidenceStatus.OFFICIAL: "高",
    EvidenceStatus.CORROBORATED: "高",
    EvidenceStatus.REFERENCE: "中",
    EvidenceStatus.INFERRED: "中",
    EvidenceStatus.PARTIAL: "低",
    EvidenceStatus.MISSING: "无",
    EvidenceStatus.MASKED: "无",
    EvidenceStatus.CONFLICT: "无",
}


class DocumentDependencyError(RuntimeError):
    """The declared ``documents`` optional dependency is unavailable."""


class DocumentComplianceError(ValueError):
    """The generated artifact failed the shared report compliance gate."""


def _require_document_dependency() -> None:
    if Document is None:
        raise DocumentDependencyError(
            "缺少 DOCX 能力（python-docx）；请安装项目 documents extra："
            "pip install 'shengxue-skill[documents]'"
        )


def _set_font(run, *, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def _style(style, *, size, color, before, after, bold=True):
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), "Calibri")
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.25


def _configure_document(document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    _style(document.styles["Heading 1"], size=16, color=HEADING_BLUE, before=18, after=10)
    _style(document.styles["Heading 2"], size=13, color=HEADING_BLUE, before=14, after=7)
    _style(document.styles["Heading 3"], size=12, color=HEADING_DARK_BLUE, before=10, after=5)

    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = Inches(1)
    section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = "匿名升学规划报告 · 可追溯证据版"
    header.paragraph_format.space_after = Pt(0)
    _set_font(header.runs[0], size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    _set_font(footer.add_run("第 "), size=8.5, color=MUTED)
    field_run = footer.add_run()
    for tag, text in (
        ("w:fldChar", None),
        ("w:instrText", " PAGE "),
        ("w:fldChar", None),
        ("w:t", "1"),
        ("w:fldChar", None),
    ):
        node = OxmlElement(tag)
        if tag == "w:fldChar":
            position = len(field_run._r)
            node.set(qn("w:fldCharType"), ("begin", "separate", "end")[(position > 0) + (position > 2)])
        else:
            node.text = text
        field_run._r.append(node)
    _set_font(field_run, size=8.5, color=MUTED)
    _set_font(footer.add_run(" 页"), size=8.5, color=MUTED)


def _numbering(document, kind: str) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (
        ("w:start", "1"),
        ("w:numFmt", kind),
        ("w:lvlText", "•" if kind == "bullet" else "%1."),
        ("w:suff", "tab"),
        ("w:lvlJc", "left"),
    ):
        node = OxmlElement(tag)
        node.set(qn("w:val"), value)
        level.append(node)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "271")
    ppr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), str(abstract_id))
    number.append(reference)
    numbering.append(number)
    return num_id


def _list_item(document, text, num_id):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend((level, number))
    paragraph._p.get_or_add_pPr().append(num_pr)
    paragraph.add_run(str(text))
    return paragraph


def _table_geometry(table, widths) -> None:
    widths = tuple(widths)
    if len(widths) != len(table.columns) or sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    table.autofit = False
    properties = table._tbl.tblPr
    for tag, value in (("w:tblW", CONTENT_WIDTH_DXA), ("w:tblInd", TABLE_INDENT_DXA)):
        node = properties.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            properties.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
    margins = OxmlElement("w:tblCellMar")
    for side, width in CELL_MARGINS_DXA.items():
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    properties.append(margins)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "C8D0DA")
        borders.append(node)
    properties.append(borders)


def _table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = str(header)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), HEADER_FILL)
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            _set_font(run, size=9.5, color=INK_BLUE, bold=True)
    header_property = OxmlElement("w:tblHeader")
    header_property.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header_property)
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = str(value)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            for run in paragraph.runs:
                _set_font(run, size=9.5)
    _table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _callout(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Pt(8)
    ppr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CALLOUT_FILL)
    ppr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:color"), HEADING_BLUE)
    borders.append(left)
    ppr.append(borders)
    _set_font(paragraph.add_run(text), color=INK_BLUE, bold=True)


def _ids(values) -> str:
    return "、".join(values) if values else "无可公开来源编号"


def _empty_recommendation_text(model):
    reason = model.recommendation_empty_reason
    if reason == "no_match_within_verified_coverage" and model.recommendation_coverage_status in {
        EvidenceStatus.OFFICIAL,
        EvidenceStatus.CORROBORATED,
        EvidenceStatus.REFERENCE,
    }:
        return "经验证覆盖范围内未找到匹配院校；未硬凑冲稳保数量。"
    return {
        "no_match_within_verified_coverage": "数据仅部分覆盖：当前已验证覆盖范围内未找到匹配院校，不能解释为没有符合院校。",
        "rank_outside_verified_coverage": "用户位次超出已验证数据覆盖范围，未生成精确推荐。",
        "unusable_evidence": "输入包含屏蔽、冲突或不可精确使用的证据，未生成数值边界。",
        "missing_verified_coverage": "缺少可验证覆盖范围，未生成精确推荐。",
    }.get(reason, "当前证据未形成可展示的普通批推荐。")


def _cover(document, model):
    # Named override: restrained report_cover derived from editorial_cover.
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(64)
    spacer.paragraph_format.space_after = Pt(0)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    _set_font(kicker.add_run("升学规划 · 证据可追溯报告"), size=10, color=HEADING_BLUE, bold=True)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    _set_font(title.add_run(f"匿名升学规划报告（{model.profile.province}）"), size=26, color=INK_BLUE, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    _set_font(
        subtitle.add_run(
            f"{model.profile.current_year} · {model.profile.subject_mode} · {model.profile.subject_selection_key}"
        ),
        size=13,
        color=MUTED,
    )
    retrieved = document.add_paragraph()
    retrieved.alignment = WD_ALIGN_PARAGRAPH.CENTER
    retrieved.paragraph_format.space_after = Pt(22)
    _set_font(retrieved.add_run(f"检索日期：{'、'.join(model.retrieval_dates)}"), size=10, color=MUTED)
    _callout(document, DISCLAIMER)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _evidence_section(document, model, bullet_id):
    profile = model.profile
    document.add_heading("一、输入与证据边界", level=1)
    rows = [
        ("年级", profile.grade),
        ("选科", f"{profile.subject_mode}；{profile.subject_selection_key}；再选科目：{_ids(profile.secondary_subjects)}"),
        ("用户提供省位次", profile.rank),
        ("能力档位", _TIER_LABEL[model.capability_tier]),
        ("查询覆盖", model.query_coverage),
        ("证据状态 / 置信度", f"{_STATUS_LABEL[model.evidence_status]} / {_CONFIDENCE_LABEL[model.evidence_status]}"),
        ("数据覆盖", _STATUS_LABEL[model.recommendation_coverage_status]),
        ("检索日期", "、".join(model.retrieval_dates)),
        ("普通批输入年份", "、".join(map(str, model.input_years)) or "无"),
        ("普通批可用年份", "、".join(map(str, model.usable_years)) or "无"),
        ("普通批策略", model.ordinary_batch_policy.policy_id),
        ("普通批策略依据", model.ordinary_batch_policy.basis_id),
        (
            "普通批检索/分档参数",
            f"检索Δ[{model.ordinary_batch_policy.search_delta_min},"
            f"{model.ordinary_batch_policy.search_delta_max}]；"
            f"冲< {model.ordinary_batch_policy.challenge_delta_lt}；"
            f"稳≤ {model.ordinary_batch_policy.stable_delta_le}；"
            f"上限冲={model.ordinary_batch_policy.tier_caps['冲']}、"
            f"稳={model.ordinary_batch_policy.tier_caps['稳']}、"
            f"保={model.ordinary_batch_policy.tier_caps['保']}",
        ),
        ("证据包标识", model.manifest_session_id),
        ("清单哈希", model.manifest_hash),
        ("来源编号", _ids(model.source_ids)),
    ]
    if model.verified_rank_coverage is not None:
        rows.append(("普通批已验证位次覆盖", f"{model.verified_rank_coverage[0]}–{model.verified_rank_coverage[1]}"))
    _table(document, ("项目", "内容"), rows, (2700, 6660))
    if model.warnings:
        document.add_heading("风险与缺失", level=2)
        for warning in model.warnings:
            _list_item(document, warning, bullet_id)


def _rank_section(document, model, bullet_id):
    document.add_heading("二、成绩定位", level=1)
    rank = model.rank
    if rank is None:
        document.add_paragraph("喜报位次证据不足：本次直接采用用户提供的省位次，不执行校排名折算。")
    elif rank.status is EvidenceStatus.INFERRED:
        for text in (
            f"推断位次区间：{rank.lower_rank}–{rank.upper_rank}",
            f"区间中位描述：{rank.median_rank} 位（推断，不是官方位次）",
            f"容差：±{rank.tolerance_rank} 位",
            f"置信度：{rank.confidence}",
            f"计算依据：{rank.method}",
            f"贡献年份：{'、'.join(map(str, rank.contributing_years))}",
            f"贡献锚点：{_ids(rank.contributing_anchor_ids)}",
            f"锚点来源编号：{_ids(rank.contributing_source_ids)}",
        ):
            _list_item(document, text, bullet_id)
    else:
        document.add_paragraph(
            f"喜报位次证据{_STATUS_LABEL[rank.status]}：{rank.reason_code or '未形成可用区间'}；未输出代理数值。"
        )


def _recommendation_section(document, model):
    document.add_heading("三、普通批冲稳保", level=1)
    if model.recommendations:
        rows = [
            (
                item.strategy,
                item.school_name,
                item.min_score,
                item.min_rank,
                _STATUS_LABEL[item.evidence_status],
                "、".join(item.source_ids),
                item.calculation_basis,
            )
            for item in model.recommendations
        ]
        _table(
            document,
            ("档位", "院校", "最低分", "最低位次", "证据状态", "来源编号", "计算依据"),
            rows,
            (600, 1500, 760, 900, 900, 1200, 3500),
        )
    else:
        document.add_paragraph(_empty_recommendation_text(model))
    for warning in model.recommendation_warnings:
        paragraph = document.add_paragraph()
        _set_font(paragraph.add_run("风险提示："), color=HEADING_DARK_BLUE, bold=True)
        paragraph.add_run(warning)
    _callout(document, DISCLAIMER)


def _pathway_section(document, model, bullet_id):
    document.add_heading("四、多元升学路径", level=1)
    if not model.pathways_available:
        document.add_paragraph("多元升学数据不足：未提供经验证的政策结果，本章节不作正式推荐。")
    elif not model.pathways:
        document.add_paragraph("多元升学数据不足：未形成正式或待核实路径；不套用无依据的位次修正。")
    else:
        status_labels = {"formal": "正式候选", "pending_verification": "待核实", "excluded": "不符合"}
        for item in model.pathways:
            document.add_heading(f"{item.title} · {item.institution}", level=2)
            for detail in (
                f"状态：{status_labels[item.status]}；资格：{item.eligibility}",
                f"政策证据状态：{_STATUS_LABEL[item.evidence_status]}；政策来源编号：{_ids(item.source_ids)}",
                f"专业选项：{_ids(item.professional_options) if item.professional_options else '当前证据未提供'}",
                f"培养安排：{item.training_arrangements or '当前证据未提供'}",
                f"转段规则：{item.transition_rules or '当前证据未提供'}",
                f"毕业/升学出口：{item.outcomes or '当前证据未提供'}",
                f"服务/就业义务：{item.service_employment_obligations or '当前证据未提供'}",
                f"退出/违约规则：{item.penalty_exit_rules or '当前证据未提供'}",
                f"费用/补助：{item.fees_and_subsidies or '当前证据未提供'}",
                f"待核实约束：{'；'.join(item.missing_constraints) or '无'}",
                f"计算依据：{item.calculation_basis}",
            ):
                _list_item(document, detail, bullet_id)
    if model.pathway_target_rank is not None:
        for detail in (
            f"有依据的路径目标位次：{model.pathway_target_rank}；位次模型证据状态：{_STATUS_LABEL[model.pathway_target_evidence_status]}",
            f"正式路径政策证据状态：{_STATUS_LABEL[model.pathway_policy_evidence_status]}",
            f"位次模型标识：{model.model_id}；方法：{model.model_method}",
            f"转换过程：{model.pathway_transformation}",
            f"模型来源编号：{_ids(model.model_source_ids)}",
        ):
            _list_item(document, detail, bullet_id)
    for warning in model.pathway_warnings:
        paragraph = document.add_paragraph()
        _set_font(paragraph.add_run("风险提示："), color=HEADING_DARK_BLUE, bold=True)
        paragraph.add_run(warning)


def _actions_section(document, model, bullet_id, decimal_id):
    document.add_heading("五、下一步行动建议", level=1)
    for action in model.action_items:
        _list_item(document, action, decimal_id)
    document.add_heading("六、证据清单与免责声明", level=1)
    document.add_paragraph("来源编号（仅展示安全编号，不展示原始 URL 或本机路径）：")
    for source_id in model.source_ids:
        _list_item(document, source_id, bullet_id)
    document.add_paragraph(f"证据包清单哈希：{model.manifest_hash}")
    _callout(document, DISCLAIMER)
    document.add_paragraph(
        "屏蔽值、冲突、部分覆盖与缺失数据均未被补成精确边界；请以省教育考试院和高校当年正式信息为准。"
    )


def _scrub(document):
    properties = document.core_properties
    properties.author = properties.last_modified_by = ""
    properties.title = "匿名升学规划报告"
    properties.subject = properties.keywords = properties.comments = ""
    properties.created = properties.modified = datetime(2000, 1, 1)


def _document_text(document):
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _save(document, output):
    source_path = None
    ready_path = None
    primary_error = None
    try:
        with tempfile.NamedTemporaryFile(
            dir=output.parent, suffix=".source.docx", delete=False
        ) as handle:
            source_path = Path(handle.name)
        with tempfile.NamedTemporaryFile(
            dir=output.parent, suffix=".ready.docx", delete=False
        ) as handle:
            ready_path = Path(handle.name)
        document.save(source_path)
        # Finish and close the deterministic ZIP under a private same-directory
        # name. The public destination cannot expose partial bytes.
        with ZipFile(source_path) as source, ZipFile(
            ready_path, "w", ZIP_DEFLATED, compresslevel=9
        ) as target:
            for name in sorted(source.namelist()):
                info = ZipInfo(name, (2000, 1, 1, 0, 0, 0))
                info.compress_type = ZIP_DEFLATED
                info.create_system = 0
                target.writestr(info, source.read(name))
        # Windows requires a writable descriptor for ``fsync`` in this runtime.
        with ready_path.open("r+b") as ready_file:
            os.fsync(ready_file.fileno())
        # A same-directory hard link publishes the already-complete inode in
        # one exclusive operation on Windows and POSIX. Existing destinations
        # raise FileExistsError and are never opened, replaced, or deleted.
        os.link(ready_path, output)
    except BaseException as error:
        primary_error = error
        raise
    finally:
        cleanup_error = None
        for temporary_path in (source_path, ready_path):
            if temporary_path is None:
                continue
            try:
                temporary_path.unlink(missing_ok=True)
            except OSError as error:
                if cleanup_error is None:
                    cleanup_error = error
        if primary_error is None and cleanup_error is not None:
            raise cleanup_error


def _output_path(model, output):
    path = Path.cwd() if output is None else Path(output)
    if path.suffix.lower() == ".docx":
        destination = path
    elif path.is_dir():
        destination = path / f"匿名升学规划报告-{model.profile.province}-{model.profile.current_year}.docx"
    else:
        raise ValueError("output must be an existing directory or a .docx path")
    destination = destination.resolve(strict=False)
    if destination.suffix.lower() != ".docx" or not destination.parent.is_dir():
        raise ValueError("DOCX output parent must exist and output must use .docx")
    if destination.exists():
        raise FileExistsError(f"refusing to overwrite existing DOCX: {destination.name}")
    return destination


def export_docx(model: ReportModel, output=None) -> Path:
    """Project one factory-authenticated model to an anonymous DOCX."""
    _require_document_dependency()
    if not isinstance(model, ReportModel):
        raise TypeError("model must be a factory-built ReportModel")
    destination = _output_path(model, output)
    document = Document()
    _configure_document(document)
    bullet_id = _numbering(document, "bullet")
    decimal_id = _numbering(document, "decimal")
    _cover(document, model)
    _evidence_section(document, model, bullet_id)
    _rank_section(document, model, bullet_id)
    _recommendation_section(document, model)
    _pathway_section(document, model, bullet_id)
    _actions_section(document, model, bullet_id, decimal_id)
    _scrub(document)
    text = _document_text(document)
    if "http://" in text or "https://" in text:
        raise DocumentComplianceError("DOCX must not contain raw URLs")
    if __package__:
        from .compliance_scan import find_price_text
    else:
        from compliance_scan import find_price_text
    hit = find_price_text(text)
    if hit is not None:
        raise DocumentComplianceError(f"DOCX compliance gate rejected: {hit}")
    _save(document, destination)
    return destination


def build_parser():
    parser = argparse.ArgumentParser(
        description="从显式省份数据、匿名画像和已验证证据包生成确定性 DOCX"
    )
    parser.add_argument("--dataset", required=True, type=Path)
    parser.add_argument("--profile", required=True, type=Path)
    parser.add_argument("--evidence", required=True, type=Path)
    parser.add_argument(
        "--secondary-subject",
        action="append",
        default=None,
        help="再选科目；可重复传入并覆盖匿名画像中的再选科目",
    )
    parser.add_argument("--output", type=Path, default=None)
    return parser


def _model_from_cli(args):
    if __package__:
        from . import generate_report as report_cli
    else:
        import generate_report as report_cli
    report_profile, recommendation_profile, pathway_profile = report_cli._load_public_profile(
        args.profile
    )
    if args.secondary_subject is not None:
        subjects = tuple(args.secondary_subject)
        report_profile = StudentProfile(
            province=report_profile.province,
            subject_mode=report_profile.subject_mode,
            subject_group=report_profile.subject_group,
            secondary_subjects=subjects,
            rank=report_profile.rank,
            grade=report_profile.grade,
            current_year=report_profile.current_year,
        )
        recommendation_profile = RecommendationProfile(
            rank=recommendation_profile.rank,
            target_province=recommendation_profile.target_province,
            subject_group=recommendation_profile.subject_group,
            secondary_subjects=frozenset(subjects),
            target_major_categories=recommendation_profile.target_major_categories,
            target_cities=recommendation_profile.target_cities,
            target_schools=recommendation_profile.target_schools,
        )
    dataset = report_cli._resolve_public_dataset(args.dataset, report_profile)
    report_profile, recommendation_profile = report_cli._profiles_with_canonical_subject_key(
        dataset, report_profile, recommendation_profile
    )
    evidence = report_cli._validated_evidence_snapshot(args.evidence)
    facts = tuple(record.to_dict() for record in evidence.facts)
    recommendations = report_cli._public_recommendations(
        dataset.admission_rows,
        recommendation_profile,
        dataset.config.ordinary_batch_policy,
        facts,
    )
    pathways = evaluate_pathways(pathway_profile, (), model=None)
    return build_report_model(
        report_profile,
        recommendations,
        rank=None,
        pathways=pathways,
        evidence=evidence,
    )


def _reconfigure_utf8():
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except (AttributeError, ValueError):
            pass


def main(argv=None):
    _reconfigure_utf8()
    try:
        _require_document_dependency()
    except DocumentDependencyError as error:
        print(f"缺少能力：{error}", file=sys.stderr)
        return 3
    args = build_parser().parse_args(argv)
    try:
        if (
            args.output is None
            or args.output.name != PUBLIC_DOCX_BASENAME
            or not args.output.parent.is_dir()
        ):
            raise ValueError("invalid public DOCX output")
        model = _model_from_cli(args)
        destination = export_docx(model, args.output)
    except (OSError, TypeError, ValueError):
        print("错误[DOCX_002]：DOCX 生成或发布失败", file=sys.stderr)
        return 2
    print(
        json.dumps(
            {
                "filename": destination.name,
                "anonymous": True,
                "secondary_subjects": list(model.profile.secondary_subjects),
            },
            ensure_ascii=False,
            sort_keys=True,
        )
    )
    return 0


__all__ = [
    "DocumentComplianceError",
    "DocumentDependencyError",
    "build_parser",
    "export_docx",
    "main",
]


if __name__ == "__main__":
    raise SystemExit(main())
